"""
DocxExtractor - extract text from Word documents.

Adapted from proven implementation in agentic/etl/transformers/extractors/docx.py.
Uses pypandoc to convert DOCX to markdown.
"""

import io
import logging
import tempfile

from agentic.ingest.extractor.base import Extractor, ExtractionError
from agentic.ingest.models import RawContent, ExtractionResult, Derivative

logger = logging.getLogger(__name__)


class DocxExtractor(Extractor):
    """
    Extract text from DOCX files, converting to markdown.
    
    Uses pypandoc for high-quality conversion that preserves structure.
    Falls back to python-docx if pypandoc is not available.
    
    Adapted from proven insurance-demo implementation.
    
    Example:
        >>> extractor = DocxExtractor()
        >>> raw = RawContent(content=docx_bytes, mime_type="application/vnd...", ...)
        >>> result = await extractor.extract(raw)
        >>> print(result.get_primary_text())
    """
    
    name = "docx"
    supported_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]
    
    async def extract(self, raw: RawContent) -> ExtractionResult:
        """
        Extract text from DOCX content.
        
        Args:
            raw: RawContent with DOCX bytes
        
        Returns:
            ExtractionResult with markdown derivative
        """
        # Try pypandoc first (preferred - better formatting)
        try:
            return self._extract_pypandoc(raw)
        except ImportError:
            logger.warning("pypandoc not available, falling back to python-docx")
        except Exception as e:
            logger.warning(f"pypandoc failed: {e}, falling back to python-docx")
        
        # Fallback to python-docx
        return self._extract_python_docx(raw)
    
    def _extract_pypandoc(self, raw: RawContent) -> ExtractionResult:
        """Extract using pypandoc - converts to markdown with good formatting."""
        try:
            import pypandoc
        except ImportError:
            raise ImportError("pypandoc not installed")
        
        # pypandoc works best with bytes directly
        markdown_text = pypandoc.convert_text(
            raw.content,
            "md",
            format="docx"
        )
        
        return ExtractionResult(
            source_uri=raw.source_uri,
            mime_type=raw.mime_type,
            derivatives=[
                Derivative(
                    type="markdown",
                    content=markdown_text,
                    format="markdown",
                ),
            ],
            auto_metadata={
                "char_count": len(markdown_text),
            },
            extraction_method="pypandoc",
            stats={"bytes_processed": len(raw.content)},
        )
    
    def _extract_python_docx(self, raw: RawContent) -> ExtractionResult:
        """Extract using python-docx as fallback."""
        try:
            from docx import Document
        except ImportError:
            raise ExtractionError(
                "Either pypandoc or python-docx is required for DOCX extraction. "
                "Install with: pip install pypandoc  OR  pip install python-docx",
                extractor_name=self.name,
                source_uri=raw.source_uri,
            )
        
        # Open document from bytes
        doc = Document(io.BytesIO(raw.content))
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Try to preserve heading structure
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else "1"
                    paragraphs.append(f"{'#' * int(level)} {text}")
                else:
                    paragraphs.append(text)
        
        # Extract tables as markdown
        for table in doc.tables:
            rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                if i == 0:  # Add header separator
                    rows.append("|" + "|".join(["---"] * len(cells)) + "|")
            if rows:
                paragraphs.append("\n".join(rows))
        
        full_text = "\n\n".join(paragraphs)
        
        # Extract document properties
        auto_metadata = {
            "char_count": len(full_text),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        
        try:
            props = doc.core_properties
            if props.title:
                auto_metadata["title"] = props.title
            if props.author:
                auto_metadata["author"] = props.author
        except Exception:
            pass
        
        return ExtractionResult(
            source_uri=raw.source_uri,
            mime_type=raw.mime_type,
            derivatives=[
                Derivative(
                    type="markdown",
                    content=full_text,
                    format="markdown",
                ),
            ],
            auto_metadata=auto_metadata,
            extraction_method="python-docx",
            stats={"bytes_processed": len(raw.content)},
        )
